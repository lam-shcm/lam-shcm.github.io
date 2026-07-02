#!/usr/bin/env python3
"""Build the formal DOCX version of 程序编写说明.md."""

from __future__ import annotations

import argparse
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "6F8F14"
ACCENT_DARK = "40520D"
CORAL = "B8462D"
INK = "202327"
MUTED = "62666D"
LIGHT = "F1F3EE"
LINE = "D8DDD0"
FONT_CN = "PingFang SC"
FONT_LATIN = "Arial"
FONT_MONO = "Menlo"


def set_run_font(run, name: str = FONT_CN, size: float | None = None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_bottom_border(paragraph, color: str = LINE, size: str = "6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, FONT_CN, 8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, FONT_CN, 8.5, color=MUTED)


def create_numbering_instance(document: Document) -> int:
    """Create a real numbered-list instance that restarts at 1."""
    numbering = document.part.numbering_part.element
    list_style = document.styles["List Number"]
    base_num_id = int(list_style._element.pPr.numPr.numId.val)

    base_num = next(
        item
        for item in numbering.findall(qn("w:num"))
        if int(item.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    new_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)

    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT_DARK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CN)
    run_properties.extend([fonts, color, underline])
    run_element.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|https?://\S+)")


def add_inline(paragraph, text: str):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True, color=INK)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, FONT_MONO, 9.2, color=CORAL)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            add_hyperlink(paragraph, token.rstrip("。，；："), token.rstrip("。，；："))
            trailing = token[len(token.rstrip("。，；：")) :]
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def configure_styles(document: Document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.42

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.35

    heading_tokens = {
        "Heading 1": (18, ACCENT_DARK, 18, 8),
        "Heading 2": (14, CORAL, 14, 6),
        "Heading 3": (11.5, INK, 10, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = FONT_MONO
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MONO)
    code_style.font.size = Pt(8.8)
    code_style.font.color.rgb = RGBColor.from_string(INK)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(5)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.2
    code_style.paragraph_format.keep_together = True

    meta_style = styles.add_style("Cover Meta", WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.name = FONT_CN
    meta_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    meta_style.font.size = Pt(11)
    meta_style.font.color.rgb = RGBColor.from_string(MUTED)
    meta_style.paragraph_format.space_after = Pt(7)


def add_cover(document: Document, student_id: str, student_name: str, school: str):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(54)

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = eyebrow.add_run("网页编程应用程序 · 本科生期末作业")
    set_run_font(run, FONT_LATIN, 10, True, ACCENT_DARK)
    eyebrow.paragraph_format.space_after = Pt(20)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("罗安民个人主页")
    set_run_font(run, FONT_CN, 30, True, INK)
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("程序编写说明")
    set_run_font(run, FONT_CN, 20, True, CORAL)
    subtitle.paragraph_format.space_after = Pt(18)
    set_bottom_border(subtitle, ACCENT, "18")

    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lead.add_run("音乐科技 × 音频工程 × 心理与认知 × AI 音乐产品")
    set_run_font(run, FONT_CN, 11, False, MUTED)
    lead.paragraph_format.space_after = Pt(48)

    for label, value in (
        ("学号", student_id),
        ("姓名", student_name),
        ("学校", school),
        ("作业方向", "选项二 · 网页编程应用程序"),
        ("在线地址", "https://lam-shcm.github.io/"),
        ("完成日期", date.today().isoformat()),
    ):
        paragraph = document.add_paragraph(style="Cover Meta")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, FONT_CN, 11, True, INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, FONT_CN, 11, False, MUTED)

    document.add_page_break()


def parse_markdown(document: Document, markdown_text: str):
    lines = markdown_text.splitlines()
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []
    numbering_id: int | None = None

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            add_inline(paragraph, " ".join(line.strip() for line in paragraph_buffer))
            paragraph_buffer = []

    def flush_code():
        nonlocal code_lines
        paragraph = document.add_paragraph(style="Code Block")
        set_cell_shading(paragraph, LIGHT)
        for index, line in enumerate(code_lines):
            run = paragraph.add_run(line)
            set_run_font(run, FONT_MONO, 8.8, color=INK)
            if index < len(code_lines) - 1:
                run.add_break()
        code_lines = []

    for line in lines:
        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph()
            numbering_id = None
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            numbering_id = None
            level = len(heading.group(1))
            if level == 1:
                continue
            text = heading.group(2)
            paragraph = document.add_heading(level=min(level - 1, 3))
            add_inline(paragraph, text)
            if level == 2:
                set_bottom_border(paragraph, LINE, "5")
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            numbering_id = None
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            if numbering_id is None:
                numbering_id = create_numbering_instance(document)
            paragraph = document.add_paragraph(style="List Number")
            apply_numbering(paragraph, numbering_id)
            add_inline(paragraph, numbered.group(1))
            continue

        paragraph_buffer.append(line)

    flush_paragraph()
    if in_code and code_lines:
        flush_code()


def build_document(
    source: Path,
    output: Path,
    student_id: str,
    student_name: str,
    school: str,
):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"{student_id} {student_name}  /  个人主页程序编写说明")
    set_run_font(run, FONT_CN, 8.5, True, MUTED)
    set_bottom_border(header, LINE, "4")
    add_page_number(section.footer.paragraphs[0])

    add_cover(document, student_id, student_name, school)
    parse_markdown(document, source.read_text(encoding="utf-8"))

    properties = document.core_properties
    properties.title = "罗安民个人主页程序编写说明"
    properties.subject = "网页编程应用程序本科期末作业"
    properties.author = student_name
    properties.last_modified_by = student_name
    properties.keywords = "HTML, CSS, JavaScript, Vite, GitHub Pages, 个人主页"
    properties.comments = f"学号：{student_id}"

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--student-id", required=True)
    parser.add_argument("--student-name", required=True)
    parser.add_argument("--school", required=True)
    args = parser.parse_args()
    build_document(
        args.source,
        args.output,
        args.student_id,
        args.student_name,
        args.school,
    )


if __name__ == "__main__":
    main()
